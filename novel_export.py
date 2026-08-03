import io
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


def _chapter_paragraphs(chapter):
    return [p for p in chapter["content"].split("\n") if p.strip()]


def _set_run_cjk_font(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def build_novel_docx(novel, chapters):
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_heading(novel["title"], level=0)
    for run in title.runs:
        _set_run_cjk_font(run)

    if novel["summary"]:
        p = doc.add_paragraph()
        run = p.add_run(novel["summary"])
        run.italic = True
        _set_run_cjk_font(run)

    prev_volume_id = "unset"
    for i, chapter in enumerate(chapters):
        if i > 0:
            doc.add_page_break()
        if chapter["volume_id"] != prev_volume_id and chapter["volume_id"]:
            vol_heading = doc.add_heading(f"第 {chapter['volume_no']} 卷 · {chapter['volume_title']}", level=0)
            for run in vol_heading.runs:
                _set_run_cjk_font(run)
        prev_volume_id = chapter["volume_id"]
        heading = doc.add_heading(f"第 {chapter['chapter_no']} 章 · {chapter['title']}", level=1)
        for run in heading.runs:
            _set_run_cjk_font(run)
        for para_text in _chapter_paragraphs(chapter):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            _set_run_cjk_font(run)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _register_pdf_font():
    # reportlab's TTFont parser is built around TrueType glyf outlines and chokes on
    # Noto Sans CJK's .ttc (it's actually a CFF-flavored OpenType collection despite the
    # extension), producing garbled text instead of raising — silently wrong, not just
    # missing. Use reportlab's built-in CID font instead: no font file is parsed at all,
    # so there's nothing to get wrong. It's a serif face, which reads fine for prose.
    name = "STSong-Light"
    pdfmetrics.registerFont(UnicodeCIDFont(name))
    return name


def build_novel_pdf(novel, chapters):
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "NovelTitle", parent=styles["Title"], fontName=font_name, fontSize=22, leading=30
    )
    summary_style = ParagraphStyle(
        "NovelSummary", parent=styles["Normal"], fontName=font_name, fontSize=10.5,
        leading=18, textColor="#666666", spaceAfter=14,
    )
    volume_style = ParagraphStyle(
        "VolumeHeading", parent=styles["Title"], fontName=font_name, fontSize=18,
        leading=26, spaceBefore=0, spaceAfter=16,
    )
    heading_style = ParagraphStyle(
        "ChapterHeading", parent=styles["Heading1"], fontName=font_name, fontSize=15,
        leading=22, spaceBefore=18, spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "ChapterBody", parent=styles["Normal"], fontName=font_name, fontSize=11,
        leading=20, spaceAfter=10, firstLineIndent=22,
    )

    story = [Paragraph(escape(novel["title"]), title_style), Spacer(1, 10)]
    if novel["summary"]:
        story.append(Paragraph(escape(novel["summary"]), summary_style))

    prev_volume_id = "unset"
    for i, chapter in enumerate(chapters):
        if i > 0:
            story.append(PageBreak())
        if chapter["volume_id"] != prev_volume_id and chapter["volume_id"]:
            story.append(Paragraph(escape(f"第 {chapter['volume_no']} 卷 · {chapter['volume_title']}"), volume_style))
        prev_volume_id = chapter["volume_id"]
        story.append(Paragraph(escape(f"第 {chapter['chapter_no']} 章 · {chapter['title']}"), heading_style))
        for para_text in _chapter_paragraphs(chapter):
            story.append(Paragraph(escape(para_text), body_style))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=2.2 * cm, bottomMargin=2.2 * cm,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm, title=novel["title"],
    )
    doc.build(story)
    buf.seek(0)
    return buf
